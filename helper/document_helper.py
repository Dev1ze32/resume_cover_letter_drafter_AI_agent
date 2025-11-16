from typing import  Optional
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from enum import Enum
from docx import Document

from helper.logger_config import get_logger

logger = get_logger(__name__)

class DocumentType(Enum):
    """Enum for document types - more maintainable than strings"""
    RESUME = "resume"
    COVER_LETTER = "cover_letter"

@dataclass
class AgentConfig:
    """Centralized configuration - easier to test and modify"""
    model_name: str = field(default_factory=lambda: os.getenv("OPENAI_MODEL", "gpt-4o-mini"))
    temperature: float = 0.5
    max_tokens: int = 500
    output_dir: Path = field(default_factory=lambda: Path("./outputs"))
    
    def __post_init__(self):
        self.output_dir.mkdir(exist_ok=True)

@dataclass
class DocumentMetadata:
    """Track document versioning and history"""
    content: str
    created_at: datetime
    last_modified: datetime
    version: int = 1
    word_count: int = 0
    
    def __post_init__(self):
        self.word_count = len(self.content.split())


class DocumentStore:
    """
    Encapsulated state management - no more globals!
    Supports versioning, history, and persistence
    """
    def __init__(self):
        self._documents: dict[DocumentType, DocumentMetadata] = {}
        self._history: list[tuple[DocumentType, str]] = []
    
    def create(self, doc_type: DocumentType, content: str) -> DocumentMetadata:
        """Create or update a document with versioning"""
        now = datetime.now()
        
        if doc_type in self._documents:
            # Update existing document
            prev = self._documents[doc_type]
            metadata = DocumentMetadata(
                content=content,
                created_at=prev.created_at,
                last_modified=now,
                version=prev.version + 1
            )
        else:
            # Create new document
            metadata = DocumentMetadata(
                content=content,
                created_at=now,
                last_modified=now
            )
        
        self._documents[doc_type] = metadata
        self._history.append((doc_type, content))
        logger.info(f"Created/updated {doc_type.value} v{metadata.version}")
        return metadata
    
    def get(self, doc_type: DocumentType) -> Optional[DocumentMetadata]:
        """Retrieve document if exists"""
        return self._documents.get(doc_type)
    
    def exists(self, doc_type: DocumentType) -> bool:
        """Check if document exists"""
        return doc_type in self._documents
    
    def get_history(self, doc_type: DocumentType) -> list[str]:
        """Get version history for a document"""
        return [content for dt, content in self._history if dt == doc_type]
    
    def clear(self):
        """Reset all documents"""
        self._documents.clear()
        self._history.clear()

    def save_to_docx(content: str, filepath: Path, doc_type: DocumentType):
        """
        Save document content to a formatted DOCX file with proper heading styles
        """
        doc = Document()
        
        # Set margins (1 inch on all sides)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Common section headings to detect
        section_keywords = [
            'summary', 'professional summary', 'experience', 'work experience',
            'education', 'skills', 'technical skills', 'certifications',
            'portfolio', 'projects', 'achievements', 'contact', 'objective',
            'career objective'
        ]
        
        # Split content into lines and add to document
        lines = content.split('\n')
        
        for line in lines:
            line_stripped = line.strip()
            
            if not line_stripped:
                # Add empty paragraph for spacing
                doc.add_paragraph()
                continue
            
            # Detect heading types
            # Check if it's a section heading (all caps single/double word like "SUMMARY", "WORK EXPERIENCE")
            is_section_all_caps = line_stripped.isupper() and 1 <= len(line_stripped.split()) <= 3
            
            # Check if it contains section keywords
            is_section_keyword = any(keyword in line_stripped.lower() for keyword in section_keywords)
            
            # Check if it's a name/title (first non-empty line, longer all caps, or has email/phone pattern)
            is_name_title = (
                line_stripped.isupper() and 
                len(line_stripped.split()) > 3 and 
                '@' not in line_stripped
            )
            
            # Check if it's a subsection (ends with colon)
            is_subsection = line_stripped.endswith(':') and len(line_stripped) < 50
            
            # Create paragraph
            paragraph = doc.add_paragraph()
            
            if is_name_title:
                # Name/Title at top - centered, bold, larger
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(line_stripped)
                run.font.name = 'Calibri'
                run.font.size = Pt(16)
                run.bold = True
                
            elif is_section_all_caps or is_section_keyword or is_subsection:
                # Section headings - bold, slightly larger
                run = paragraph.add_run(line_stripped)
                run.font.name = 'Calibri'
                run.font.size = Pt(12)
                run.bold = True
                # Add spacing before section
                paragraph.paragraph_format.space_before = Pt(12)
                paragraph.paragraph_format.space_after = Pt(6)
                
            else:
                # Regular content
                run = paragraph.add_run(line_stripped)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
                
                # Handle bullet points
                if line_stripped.startswith('•') or line_stripped.startswith('-'):
                    paragraph.style = 'List Bullet'
        
        # Save the document
        doc.save(str(filepath))
        logger.info(f"Saved DOCX to {filepath}")
